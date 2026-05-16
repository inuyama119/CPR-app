import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # Title
    title = doc.add_heading('胸骨圧迫マスター アプリケーション概要資料', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 1
    doc.add_heading('1. アプリケーションの目的', level=1)
    doc.add_paragraph('本アプリケーションは、「市民が救命講習を受講した後の自主的な復習」を支援するために開発されました。')
    doc.add_paragraph('救命講習で学んだ「胸骨圧迫（心臓マッサージ）」の技術は、時間の経過とともに忘れてしまいがちです。本アプリを使用することで、いつでもどこでもスマートフォン1つで、正しいリズムと姿勢の自己評価が可能になります。')

    # Section 2
    doc.add_heading('2. 主な機能と使い方', level=1)
    doc.add_paragraph('本アプリには、手軽に練習できる「姿勢復習モード」と、より厳密に評価を行う「ビデオ精密評価モード」の2つの機能があります。')
    
    doc.add_heading('① 姿勢復習モード（手軽な自己練習）', level=2)
    doc.add_paragraph('スマートフォンの画面に表示されるガイドイラストと、最適なテンポ（110BPM）を刻むリズム音に合わせて胸骨圧迫の練習を行います。\n直感的なUIで、正しい体重のかけ方や腕の角度を視覚的に復習できます。')

    doc.add_heading('② ビデオ精密評価モード（AIによる本格的な動作解析）', level=2)
    doc.add_paragraph('ユーザーが自身の胸骨圧迫を真横から撮影した動画（約20秒）を読み込ませることで、Googleの最新AIモデル（MediaPipe Pose）が骨格を自動検出し、圧迫の「リズム」と「腕の垂直度」を精密に評価します。')
    
    p = doc.add_paragraph()
    p.add_run('【ビデオ評価の手順】\n').bold = True
    p.add_run('1. トップ画面から「ビデオで精密評価」を選択。\n')
    p.add_run('2. 事前に撮影した胸骨圧迫の動画（真横から全身を写したもの）を選択。\n')
    p.add_run('3. AIによる骨格解析がスタート。\n')
    p.add_run('4. 解析完了後、詳細な診断レポートが表示されます。')

    # Section 3
    doc.add_heading('3. 診断レポート（UIパターンの解説）', level=1)
    doc.add_paragraph('AIによる解析結果は、ユーザーのモチベーションを下げないよう、ポジティブで分かりやすい3段階の「級」で表示されます。姿勢やリズムが基準に満たない場合は、具体的なアドバイスが提示されます。')

    # Pattern A
    doc.add_heading('パターンA：エキスパート級（すべての基準をクリア）', level=2)
    doc.add_paragraph('リズム（100〜120回/分）と姿勢（腕の傾き20度以内）の両方が完璧な場合の表示です。\n達成感を高めるため、ゴールドのカラーリングと紙吹雪のアニメーションで最高評価を演出します。')
    try:
        doc.add_picture('expert_evaluation.png', width=Inches(3.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f'[画像挿入エラー: {e}]')

    # Pattern B
    doc.add_heading('パターンB：スタンダード級（どちらか1つが基準外）', level=2)
    doc.add_paragraph('リズムまたは姿勢のどちらか一方が惜しくも基準に満たなかった場合の表示です。\n「素晴らしい技術をお持ちです」と肯定しつつ、不足していた項目に対する具体的なアドバイス（例：「音をよく聞いて練習しましょう」）を表示し、次のステップアップを促します。')
    try:
        doc.add_picture('standard_evaluation.png', width=Inches(3.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f'[画像挿入エラー: {e}]')

    # Pattern C
    doc.add_heading('パターンC：あと一歩！（両方が基準外）', level=2)
    p = doc.add_paragraph('リズムと姿勢の')
    p.add_run('両方が基準に満たなかった場合').bold = True
    p.add_run('の表示です。\n「トレーニングが必要」「不合格」といったネガティブな表現を避け、')
    p.add_run('「まずは練習に挑戦したことが素晴らしい一歩です」').bold = True
    p.add_run('とユーザーの行動自体を称賛します。その上で、どこを直せば良いのか（リズムと姿勢の両方のアドバイス）を優しく提示します。')
    try:
        doc.add_picture('training_evaluation.png', width=Inches(3.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f'[画像挿入エラー: {e}]')

    # Section 4
    doc.add_heading('4. 期待される効果', level=1)
    doc.add_paragraph('・技術維持率の向上: 定期的な自己評価により、講習内容の忘却を防ぎます。\n・心理的ハードルの低下: いつでもゲーム感覚で評価を受けられるため、トレーニングへの参加意欲が高まります。\n・質の高い救命処置の普及: 正しい姿勢とリズムが身につくことで、実際の救急現場での救命率向上に貢献します。')

    doc.save('CPR_App_Presentation.docx')
    print("Word document created successfully.")

if __name__ == "__main__":
    main()
